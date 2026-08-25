#!/usr/bin/env python3
"""Validate the reference apparatus of the Human Mutation manuscript/supplementary.

The bundled citation-management validator expects BibTeX; these documents carry
plain-text Vancouver references in .docx, so this checks what actually matters
here:

  1. every in-text marker resolves to a reference entry
  2. every reference entry is cited at least once
  3. reference numbering is contiguous from 1
  4. each entry's DOI (where present) resolves on Crossref and its title matches
  5. entries with no DOI are looked up by title so a DOI can be added

The two documents use different marker conventions: the manuscript uses runs
with font.superscript = True carrying plain digits; the supplementary uses
Unicode superscript characters inline. Both are handled.

Usage:
  python3 analyses/validate_manuscript_citations.py [--check-dois]
"""

import argparse
import difflib
import json
import re
import sys
import time
import urllib.parse
import urllib.request

import docx

DOCS = [
    ("manuscript", "docs/Metpally_VarViz_HumMutat_Manuscript.docx", "brackets"),
    ("supplementary", "docs/Metpally_VarViz_HumMutat_Supplementary.docx", "brackets"),
]

# Journal style is "[9]" / "[9, 10]" / "[14-20]". Brackets are also used for
# CRediT roles in Author Contributions ("[equal]") and for the supplementary's
# author line, so a bracket only counts when it holds digits and separators.
BRACKET_RE = re.compile(r"\[([\d][\d,;\s\u2013\u2014-]*)\]")

UNSUP = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹", "0123456789")
SUPCHARS = "⁰¹²³⁴⁵⁶⁷⁸⁹"
SUPMINUS = "\u207b"  # U+207B SUPERSCRIPT MINUS, the exponent tell
SEPS = ",;\u02d2"    # U+02D2 MODIFIER LETTER CENTRED RIGHT HALF RING
MARKER_RE = re.compile(f"[{SUPCHARS}][{SUPCHARS}{SEPS}–—-]*")
DOI_RE = re.compile(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+")


def unicode_tokens(text):
    """Citation markers written as Unicode superscript characters.

    Exponents are indistinguishable from markers by shape: "1×10⁻⁴" ends in a
    superscript 4, "4.17 × 10⁻⁵" in a 5. Counting them inflates the cited set
    with whatever numbers the exponents happen to be -- harmless while those
    land inside the reference range, then silently wrong.

    Every exponent in these documents is preceded by superscript minus and no
    genuine marker is, so that is the discriminator. It cannot be a
    digit-predecessor test: real markers do follow digits, as in a parameter
    value or a database version immediately before its citation.
    """
    for m in MARKER_RE.finditer(text):
        if m.start() and text[m.start() - 1] == SUPMINUS:
            continue
        # the class admits a trailing separator, as in a marker before a comma
        tok = m.group(0).rstrip(SEPS + "–— ").translate(UNSUP)
        if re.search(r"\d", tok):
            yield tok


def _selftest():
    """Synthetic fixtures, not lines from the documents: this repository is
    public and the manuscript is not."""
    got = list(unicode_tokens("threshold 1×10⁻⁴ and cutoff 5×10⁻² apply"))
    assert got == [], f"exponents leaked through: {got}"
    assert list(unicode_tokens("parameter 0.2¹, next clause")) == ["1"]   # marker after a digit
    assert list(unicode_tokens("release v4.1²².")) == ["22"]              # and after a version
    assert list(unicode_tokens("two tools ¹¹˒¹²")) == ["11˒12"]           # U+02D2 separator
    assert list(unicode_tokens("sources²⁸,²⁹, then more")) == ["28,29"]   # comma pair
    assert list(unicode_tokens("4.17 × 10⁻⁵ alongside 0.2¹")) == ["1"]    # both in one line
    print("self-test OK")


def norm_title(s):
    return re.sub(r"[^a-z0-9 ]", "", s.lower()).strip()


def crossref(path):
    url = "https://api.crossref.org/" + path
    req = urllib.request.Request(url, headers={"User-Agent": "varviz-citation-check"})
    with urllib.request.urlopen(req, timeout=45) as fh:
        return json.load(fh)["message"]


def extract(doc_path, mode):
    d = docx.Document(doc_path)
    paras = d.paragraphs
    iref = next(i for i, p in enumerate(paras) if p.text.strip() == "References")
    end = next(
        (
            i
            for i, p in enumerate(paras)
            if p.text.strip() == "Figure Legends" and i > iref
        ),
        len(paras),
    )
    refs = [p.text.strip() for p in paras[iref + 1 : end] if p.text.strip()]

    def expand(s):
        """'14-20' -> 14..20, '32,33' -> [32, 33]. Ranges are written with a
        hyphen or an en dash; reading them as two endpoints silently reports
        every reference inside the range as uncited."""
        out = []
        for part in re.split(r"[,;]", s):
            part = part.strip()
            m = re.fullmatch(r"(\d+)\s*[-–—]\s*(\d+)", part)
            if m:
                a, b = int(m.group(1)), int(m.group(2))
                out += list(range(min(a, b), max(a, b) + 1))
            else:
                out += [int(x) for x in re.findall(r"\d+", part)]
        return out

    def markers(p):
        out = []
        if mode == "brackets":
            for m in BRACKET_RE.finditer(p.text):
                out += expand(m.group(1))
        elif mode == "runs":
            for r in p.runs:
                if r.font.superscript and r.text.strip():
                    out += expand(r.text)
        else:
            out += [t for tok in unicode_tokens(p.text) for t in expand(tok)]
        return out

    cited = []
    for i, p in enumerate(paras):
        if i > iref:
            break
        cited += markers(p)

    # Tool-comparison and data-source tables carry citations of their own, and
    # they live outside document.paragraphs -- missing them reports every
    # table-only reference as uncited.
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    cited += markers(p)
    return refs, cited


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--check-dois", action="store_true", help="verify against Crossref (slow)"
    )
    ap.add_argument(
        "--self-test", action="store_true", help="check marker parsing and exit"
    )
    args = ap.parse_args()

    if args.self_test:
        _selftest()
        return 0

    problems = 0
    for name, path, mode in DOCS:
        print(f"\n{'=' * 70}\n{name.upper()}  ({path})\n{'=' * 70}")
        refs, cited = extract(path, mode)
        n = len(refs)
        cset = set(cited)
        print(f"  references: {n}    distinct in-text markers: {len(cset)}")

        dangling = sorted(x for x in cset if x < 1 or x > n)
        if dangling:
            problems += 1
            print(f"  ✗ markers with no reference entry: {dangling}")
        uncited = sorted(set(range(1, n + 1)) - cset)
        if uncited:
            problems += 1
            print(f"  ✗ reference entries never cited: {uncited}")
            for u in uncited:
                print(f"      {u}. {refs[u - 1][:95]}")
        if not dangling and not uncited:
            print("  ✓ markers and entries correspond exactly")

        if not args.check_dois:
            continue

        print("\n  Crossref verification:")
        for i, ref in enumerate(refs, 1):
            m = DOI_RE.search(ref)
            title = re.split(r"\.\s", ref, maxsplit=1)
            title = title[1] if len(title) > 1 else ref
            title = title.split(".")[0].strip()
            try:
                if m:
                    doi = m.group(0).rstrip(".")
                    msg = crossref("works/" + urllib.parse.quote(doi))
                    ct = msg.get("title", [""])[0]
                    ratio = difflib.SequenceMatcher(
                        None, norm_title(title), norm_title(ct)
                    ).ratio()
                    flag = "✓" if ratio > 0.75 else "✗ TITLE MISMATCH"
                    if ratio <= 0.75:
                        problems += 1
                    print(f"   {flag} {i:2d}. {doi}")
                    if ratio <= 0.75:
                        print(f"          doc: {title[:80]}")
                        print(f"          doi: {ct[:80]}")
                else:
                    q = crossref(
                        "works?rows=1&query.bibliographic="
                        + urllib.parse.quote(title[:180])
                    )
                    it = q.get("items", [])
                    if it:
                        ct = it[0].get("title", [""])[0]
                        ratio = difflib.SequenceMatcher(
                            None, norm_title(title), norm_title(ct)
                        ).ratio()
                        if ratio > 0.85:
                            print(
                                f"   · {i:2d}. no DOI in entry; Crossref suggests "
                                f"{it[0].get('DOI')}"
                            )
                        else:
                            print(
                                f"   ? {i:2d}. no DOI, no confident Crossref match "
                                f"({title[:60]})"
                            )
                    else:
                        print(f"   ? {i:2d}. no DOI, no Crossref result ({title[:60]})")
            except Exception as e:
                print(f"   ! {i:2d}. lookup failed: {repr(e)[:70]}")
            time.sleep(0.4)

    print(f"\n{'=' * 70}")
    print(
        "PASS - no structural problems"
        if not problems
        else f"{problems} problem group(s) above"
    )
    return 1 if problems else 0


if __name__ == "__main__":
    sys.exit(main())
